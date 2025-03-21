#!/usr/bin/env python 3
from docx.shared import Inches
from app.util.globals import flex
import xlsxwriter
import docx

from app.util.courses import Course
from app.util.students import Student

def putMasterTimetable(
	table:		  dict[str: dict], 
	output_dir: str = './output/final'
) -> None:

	workbook = xlsxwriter.Workbook(f'{output_dir}/master_timetable.xlsx')
	worksheet = workbook.add_worksheet()

	sem1Header = workbook.add_format({'bold': True, 'align': 'center', 'bg_color': '#91c5fa', 'border': 1})
	sem2Header = workbook.add_format({'bold': True, 'align': 'center', 'bg_color': '#a2fa91', 'border': 1})
	sem1Data = workbook.add_format({'bg_color': '#cee1f5', 'border': 1})
	sem2Data = workbook.add_format({'bg_color': '#d8f7d2', 'border': 1})
	worksheet.merge_range(0, 0, 0, len(table) * 2 - 1, 'Semester 1', sem1Header)
	worksheet.merge_range(0, len(table) * 2, 0, (len(table) * 4) - 1, 'Semester 2', sem2Header)

	tracker = 0
	for i in range(0, len(table) * 4, 4):
		tracker += 1
		block, index = tracker, f'block{tracker}'
		header = sem1Header if tracker <= len(table) / 2 else sem2Header	
		block = block if block <= len(table) / 2 else block - (len(table) // 2)
		worksheet.merge_range(1, i, 1, i + 3, f'Block {block}', header)

		row = 2
		for course in table[index]:
			style = sem1Data if tracker <= len(table) / 2 else sem2Data
			worksheet.merge_range(row, i, row + 1, i + 3, f'{table[index][course]["Description"]}\n{course[:-2]}', style)
			row += 2

	workbook.close()

def putScheduleToWord(
	courses: 		dict[str: Course], 
	student: 		Student, 
	output_dir: str = './output/final/student_schedules'
) -> None:

	# create an instance of a word doc.
	doc = docx.Document()

	doc.add_heading(f'Schedule for Student {student.PupilNum}')

	# Create a table object
	table = doc.add_table(rows=1, cols=3)
	table.autofit = True
	table.style = 'Colorful List'

	# course name + (code) | block | semester
	row = table.rows[0].cells
	row[0].text = 'Course'
	row[1].text = 'Block'
	row[2].text = 'Sem.'
 
	table.columns[0].width = Inches(3.5)
	table.columns[1].width = Inches(0.75)
	table.columns[2].width = Inches(0.75)

	for block in student.Schedule:
		courseCode = student.Schedule[block][0]
		if courseCode in flex: courseName = 'Study'
		else: courseName = courses[courseCode[:-2]].Description
		blockNum = int(block.split('block')[1])
		semester = 2 if blockNum > 5 else 1
		if blockNum > (len(student.Schedule)/2): blockNum -= (len(student.Schedule)/2)
		row = table.add_row().cells
		row[0].text = f'{courseName}\n({courseCode})'
		row[1].text = str(blockNum)
		row[2].text = str(semester)

	for row in table.rows:
		row.height = Inches(0.75)

	doc.save(f'{output_dir}/{student.PupilNum}_schedule.docx')
